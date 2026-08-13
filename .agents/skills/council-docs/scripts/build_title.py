"""Build the GOST PhD dissertation TITLE PAGE (EN + KZ) as .docx (+ .pdf).

The title page is positional (centered organization block; UDC left / "as a
manuscript" right on one line; centered author, title, programme, degree
statement; a consultant block; centered place/year at the bottom), so it is
built directly with python-docx rather than through the Markdown converter. It
reuses md2gost's GOST page/style/font helpers so the page geometry (A4, TNR 14,
margins 30/10/20/20 mm) matches every other deliverable.

Required elements follow council/en/10-dissertation/structure.md §3.1.1:
organization, UDC, full name, title, programme code+name, sought degree,
scientific consultant, place/year. Their values are read from the metadata
registry council/METADATA.toml, so the title page cannot drift away from the
abstracts and the reviews (it once did: it carried a superseded Kazakh title).

Usage:
    python build_title.py [--date YYYY-MM-DD] [--no-pdf]
"""
from __future__ import annotations

import argparse
import importlib.util
import re
import tomllib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Mm, Pt

_HERE = Path(__file__).resolve().parent
ROOT = _HERE
while ROOT.parent != ROOT and not (ROOT / "defense").is_dir():
    ROOT = ROOT.parent

_spec = importlib.util.spec_from_file_location("md2gost", _HERE / "md2gost.py")
md2gost = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(md2gost)

TAB_MM = 170.0  # text-area right edge: A4 210 - left 30 - right 10

# --- field values -------------------------------------------------------------
# Every name, code and title on the page comes from the single registry
# council/METADATA.toml. Only the layout labels ("Scientific consultant", the
# degree statement, "On manuscript right") live here, because they are part of
# the GOST form rather than data about this defense.


def _registry() -> dict:
    with (ROOT / "council/METADATA.toml").open("rb") as fh:
        return tomllib.load(fh)


def _fields(reg: dict) -> dict:
    org, prog, diss = reg["organization"], reg["programme"], reg["dissertation"]
    sup, fc = reg["supervisor"], reg["foreign_consultant"]
    year = diss["year"]
    return {
        "en": {
            "org": [org["name_en"]],
            "udc": f"UDC: {diss['udc']}",
            "manuscript": "On manuscript right",
            "author": reg["candidate"]["name_upper_en"],
            "title": diss["title_en"],
            "programme": f"{prog['code']} – {prog['name_en']}",
            "degree": [
                "Thesis for the degree of doctor of",
                "Philosophy (PhD)",
            ],
            "consultant": [
                ("Scientific consultant", False),
                (f"{sup['degree_en_short']},", False),
                (f"{sup['title_en']}, {sup['org_en']}", False),
                (sup["short_en"], False),
                ("", False),
                ("Foreign consultant", False),
                (f"{fc['title_en']}, {fc['org_en']}", False),
                (fc["short_en"], False),
            ],
            "place": [org["country_en"], f"{org['city_en']}, {year}"],
        },
        "kz": {
            "org": [org["name_kz"]],
            "udc": f"ӘОЖ: {diss['udc']}",
            "manuscript": "Қолжазба құқығында",
            "author": reg["candidate"]["name_upper_kz"],
            "title": diss["title_kz"],
            "programme": f"{prog['code']} – {prog['name_kz']}",
            "degree": [
                "Философия докторы (PhD) дәрежесін",
                "алуға арналған диссертация",
            ],
            "consultant": [
                ("Ғылыми консультанты", False),
                (f"{sup['degree_kz']},", False),
                (f"{sup['title_kz']}, {sup['org_kz']}", False),
                (sup["short_kz"], False),
                ("", False),
                ("Шетелдік ғылыми консультанты", False),
                (f"{fc['title_kz']}, {fc['org_en']}", False),
                (fc["short_en"], False),
            ],
            "place": [org["country_kz"], f"Алматы, {year}"],
        },
    }


FIELDS = _fields(_registry())


def _centered(doc, text, *, bold=False, size=14, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Mm(0)
    pf.line_spacing = 1.0
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    r = p.add_run(text)
    md2gost._set_cell_font(r, bold=bold)
    r.font.size = Pt(size)
    return p


def _gap(doc, n):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Mm(0)


def populate(doc, lang: str) -> None:
    """Add the title-page content to an existing (already-configured) document.

    Used both by `build()` (standalone title page) and by the front-matter
    bundle, which composes the title page as the first page of one document.
    """
    f = FIELDS[lang]

    # Organization block (top, centered, bold)
    for line in f["org"]:
        _centered(doc, line, bold=True, size=14, space_after=2)

    # UDC (left) / "as a manuscript" (right) on one line
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(TAB_MM), WD_TAB_ALIGNMENT.RIGHT)
    md2gost._set_cell_font(p.add_run(f["udc"]))
    md2gost._set_cell_font(p.add_run("\t" + f["manuscript"]))

    _gap(doc, 7)
    _centered(doc, f["author"], bold=True, size=14)

    _gap(doc, 4)
    _centered(doc, f["title"], bold=True, size=16, space_after=6)
    _centered(doc, f["programme"], bold=False, size=14, space_before=12)

    _gap(doc, 1)
    for line in f["degree"]:
        _centered(doc, line, bold=False, size=14)

    # Consultant block — left-aligned, shifted into the right half of the page
    _gap(doc, 3)
    for text, bold in f["consultant"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.left_indent = Mm(85)
        pf.first_line_indent = Mm(0)
        pf.line_spacing = 1.0
        pf.space_after = Pt(0)
        if text:
            md2gost._set_cell_font(p.add_run(text), bold=bold)

    # Place / year — centered, near the bottom
    _gap(doc, 6)
    for line in f["place"]:
        _centered(doc, line, bold=False, size=14)


def build(lang: str, out_docx: Path) -> None:
    doc = Document()
    md2gost._configure_styles(doc)
    md2gost._configure_page(doc)
    populate(doc, lang)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    print("[docx]", out_docx.name)


def latest_manuscript_date() -> str:
    """Date stamp of the newest rendered manuscript pair in defense/docs/.

    The title page carries no dated input — the stamp is a label only — but a
    pinned label is worse than no label: it is how the other four builders sent
    current content out under a June name, and it is why this one produced no
    August file at all while everything beside it was rebuilt.
    """
    docs = ROOT / "defense/docs"
    dates = sorted(
        m.group(1)
        for p in docs.glob("DISSERTATION_EN_GOST_*.docx")
        if (m := re.search(r"_(\d{4}-\d{2}-\d{2})\.docx$", p.name))
        and (docs / p.name.replace("_EN_", "_KZ_")).is_file()
    )
    if not dates:
        raise SystemExit(f"no DISSERTATION_{{EN,KZ}}_GOST_*.docx pair in {docs}")
    return dates[-1]


def main() -> None:
    ap = argparse.ArgumentParser(description="Build GOST TITLE PAGE (EN+KZ)")
    ap.add_argument("--date", default=None, help="output date stamp (default: newest manuscript)")
    ap.add_argument("--no-pdf", action="store_true")
    args = ap.parse_args()
    if args.date is None:
        args.date = latest_manuscript_date()
        print(f"[src ] newest manuscript: {args.date}")

    out_dir = ROOT / "defense/docs/front_matter"
    built = []
    for lang in ("en", "kz"):
        out = out_dir / f"TITLE_PAGE_{lang.upper()}_GOST_{args.date}.docx"
        build(lang, out)
        built.append(out)

    if not args.no_pdf:
        from docx2pdf import convert
        for out_docx in built:
            pdf = out_docx.with_suffix(".pdf")
            convert(str(out_docx), str(pdf))
            print("[pdf ]", pdf.name)


if __name__ == "__main__":
    main()
