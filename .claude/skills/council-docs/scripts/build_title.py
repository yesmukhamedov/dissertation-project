"""Build the GOST PhD dissertation TITLE PAGE (EN + KZ) as .docx (+ .pdf).

The title page is positional (centered organization block; UDC left / "as a
manuscript" right on one line; centered author, title, programme, degree
statement; a consultant block; centered place/year at the bottom), so it is
built directly with python-docx rather than through the Markdown converter. It
reuses md2gost's GOST page/style/font helpers so the page geometry (A4, TNR 14,
margins 30/10/20/20 mm) matches every other deliverable.

Required elements follow council/en/10-dissertation/structure.md §3.1.1:
organization, UDC, full name, title, programme code+name, sought degree,
scientific consultant, place/year. Their values are read from the metadata
registry council/METADATA.toml, so the title page cannot drift away from the
abstracts and the reviews (it once did: it carried a superseded Kazakh title).

Usage:
    python build_title.py [--date YYYY-MM-DD] [--no-pdf]
"""
from __future__ import annotations

import argparse
import importlib.util
import os
import re
import tomllib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Mm, Pt

_HERE = Path(__file__).resolve().parent
ROOT = _HERE
while ROOT.parent != ROOT and not (ROOT / "defense").is_dir():
    ROOT = ROOT.parent

_spec = importlib.util.spec_from_file_location("md2gost", _HERE / "md2gost.py")
md2gost = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(md2gost)

TAB_MM = 170.0  # text-area right edge: A4 210 - left 30 - right 10

# --- field values -------------------------------------------------------------
# Every name, code and title on the page comes from the single registry
# council/METADATA.toml. Only the layout labels ("Scientific consultant", the
# degree statement, "On manuscript right") live here, because they are part of
# the GOST form rather than data about this defense.


def _registry() -> dict:
    with (ROOT / "council/METADATA.toml").open("rb") as fh:
        return tomllib.load(fh)


def _fields(reg: dict) -> dict:
    org, prog, diss = reg["organization"], reg["programme"], reg["dissertation"]
    sup, fc = reg["supervisor"], reg["foreign_consultant"]
    year = diss["year"]
    return {
        "en": {
            "org": [org["name_en"]],
            "udc": f"UDC: {diss['udc']}",
            "manuscript": "On manuscript right",
            "author": reg["candidate"]["name_upper_en"],
            "title": diss["title_en"],
            "programme": f"{prog['code']} – {prog['name_en']}",
            "degree": [
                "Thesis for the degree of doctor of",
                "Philosophy (PhD)",
            ],
            "consultant": [
                ("Scientific consultant", False),
                (f"{sup['degree_en_short']},", False),
                (f"{sup['title_en']}, {sup['org_en']}", False),
                (sup["short_en"], False),
                ("", False),
                ("Foreign consultant", False),
                (f"{fc['title_en']}, {fc['org_en']}", False),
                (fc["short_en"], False),
            ],
            "place": [org["country_en"], f"{org['city_en']}, {year}"],
        },
        "kz": {
            "org": [org["name_kz"]],
            "udc": f"ӘОЖ: {diss['udc']}",
            "manuscript": "Қолжазба құқығында",
            "author": reg["candidate"]["name_upper_kz"],
            "title": diss["title_kz"],
            "programme": f"{prog['code']} – {prog['name_kz']}",
            "degree": [
                "Философия докторы (PhD) дәрежесін",
                "алуға арналған диссертация",
            ],
            "consultant": [
                ("Ғылыми консультанты", False),
                (f"{sup['degree_kz']},", False),
                (f"{sup['title_kz']}, {sup['org_kz']}", False),
                (sup["short_kz"], False),
                ("", False),
                ("Шетелдік ғылыми консультанты", False),
                (f"{fc['title_kz']}, {fc['org_en']}", False),
                (fc["short_en"], False),
            ],
            "place": [org["country_kz"], f"Алматы, {year}"],
        },
    }


FIELDS = _fields(_registry())


# --- one-page fitting ---------------------------------------------------------
# The frame is fixed (A4 minus the GOST margins) but the content is data: the
# Kazakh title wraps to three lines where the English one takes two, and the
# consultant block wraps differently again in each language. Hard-coded blank-line
# gaps could not hold that — they pushed "Алматы, 2026" off the Kazakh title page
# onto a second page. So the gaps are computed: every line is measured, the
# leftover height is handed to the gaps, and the place/year block is pushed to the
# bottom of the page, the way it sits in the council sample title pages.

PT_PER_MM = 72.0 / 25.4
BODY_H_PT = (297 - 20 - 20) * PT_PER_MM   # A4 height minus top/bottom margins
TEXT_W_PT = 170.0 * PT_PER_MM             # text column: 210 - 30 - 10
CONSULT_INDENT_MM = 85.0                  # consultant block sits in the right half
LINE_FACTOR = 1.15                        # Word single spacing, Times New Roman
BODY_PT = 14                              # Normal style size
GAP_PT = BODY_PT * LINE_FACTOR            # height of one blank paragraph
# One blank line held back from the budget: the wrap is measured here from the
# Times New Roman metrics, and Word is entitled to break one line differently.
# A spare line absorbs exactly that, and the slack costs nothing on the page.
SAFETY_PT = GAP_PT

# Preferred blank-line gaps, in source order:
#   0 after the UDC line, 1 before the title, 2 before the degree statement,
#   3 before the consultant block, 4 before place/year.
# Gap 4 is elastic — it takes whatever height is left over — and the rest shrink
# proportionally only when the page cannot hold them.
GAPS = (7, 4, 1, 3, 6)

_FONTS = {False: "times.ttf", True: "timesbd.ttf"}
_font_cache: dict[tuple[str, int], object] = {}


def _font(size_pt: int, bold: bool):
    """Times New Roman at 4× scale (integer sizes measure more precisely)."""
    key = (_FONTS[bold], size_pt)
    if key not in _font_cache:
        from PIL import ImageFont
        _font_cache[key] = ImageFont.truetype(
            str(Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / _FONTS[bold]),
            size_pt * 4,
        )
    return _font_cache[key]


def _wrapped_lines(text: str, size_pt: int, bold: bool, width_pt: float) -> int:
    """Number of lines `text` occupies in a `width_pt`-wide column (greedy wrap)."""
    words = text.split()
    if not words:
        return 1
    f = _font(size_pt, bold)
    space = f.getlength(" ") / 4.0
    lines, cur = 1, f.getlength(words[0]) / 4.0
    for w in words[1:]:
        ww = f.getlength(w) / 4.0
        if cur + space + ww > width_pt:
            lines, cur = lines + 1, ww
        else:
            cur += space + ww
    return lines


def _para(text, *, bold=False, size=BODY_PT, align="center", before=0, after=0,
          indent_mm=0.0, tab=None) -> dict:
    return {"text": text, "bold": bold, "size": size, "align": align,
            "before": before, "after": after, "indent_mm": indent_mm, "tab": tab}


def _layout(lang: str) -> list[dict]:
    """The title page as an ordered list of paragraphs and gap slots.

    One description, consumed twice — measured by `_solve_gaps` and rendered by
    `populate` — so the measurement can never drift from what is written out.
    """
    f = FIELDS[lang]
    items: list[dict] = []

    # Organization block (top, centered). Not bold: the council sample title
    # pages carry it in the regular face, and the university name is not a
    # heading here.
    items += [_para(line, after=2) for line in f["org"]]

    # UDC (left) / "as a manuscript" (right) on one line
    items.append(_para(f["udc"], align="left", before=18, tab=f["manuscript"]))

    items.append({"gap": 0})
    items.append(_para(f["author"], bold=True))

    items.append({"gap": 1})
    items.append(_para(f["title"], bold=True, size=16, after=6))
    items.append(_para(f["programme"], before=12))

    items.append({"gap": 2})
    items += [_para(line) for line in f["degree"]]

    items.append({"gap": 3})
    # Flush right against the text-area edge, the way the council sample title
    # pages set the consultant block; the left indent only bounds the wrap so it
    # keeps to the right half of the page.
    items += [_para(text, bold=bold, align="right", indent_mm=CONSULT_INDENT_MM)
              for text, bold in f["consultant"]]

    items.append({"gap": 4})
    items += [_para(line) for line in f["place"]]
    return items


def _height(item: dict) -> float:
    width = TEXT_W_PT - item["indent_mm"] * PT_PER_MM
    lines = _wrapped_lines(item["text"], item["size"], item["bold"], width)
    return lines * item["size"] * LINE_FACTOR + item["before"] + item["after"]


def _solve_gaps(items: list[dict]) -> list[int]:
    """Blank lines per gap slot so the whole page fits on one page.

    With room to spare the last gap absorbs the slack, seating place/year at the
    foot of the page; when the text is too tall for the preferred gaps, all five
    shrink proportionally (largest remainder), never below zero.
    """
    fixed = sum(_height(i) for i in items if "gap" not in i)
    budget = int((BODY_H_PT - SAFETY_PT - fixed) // GAP_PT)
    if budget <= 0:
        return [0] * len(GAPS)
    head = sum(GAPS[:-1])
    if budget > head:
        return [*GAPS[:-1], budget - head]
    # Proportional shrink, largest remainder first.
    total = sum(GAPS)
    exact = [g * budget / total for g in GAPS]
    gaps = [int(e) for e in exact]
    for idx in sorted(range(len(GAPS)), key=lambda i: exact[i] - gaps[i], reverse=True):
        if sum(gaps) >= budget:
            break
        gaps[idx] += 1
    return gaps


def _emit(doc, item: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(item["align"], WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.first_line_indent = Mm(0)
    pf.left_indent = Mm(item["indent_mm"])
    pf.line_spacing = 1.0
    pf.space_before = Pt(item["before"])
    pf.space_after = Pt(item["after"])
    if item["tab"] is not None:
        pf.tab_stops.add_tab_stop(Mm(TAB_MM), WD_TAB_ALIGNMENT.RIGHT)
    if item["text"]:
        r = p.add_run(item["text"])
        md2gost._set_cell_font(r, bold=item["bold"])
        r.font.size = Pt(item["size"])
    if item["tab"] is not None:
        md2gost._set_cell_font(p.add_run("	" + item["tab"]))


def _gap(doc, n: int) -> None:
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Mm(0)


def populate(doc, lang: str) -> None:
    """Add the title-page content to an existing (already-configured) document.

    Used both by `build()` (standalone title page) and by the front-matter
    bundle, which composes the title page as the first page of one document.
    """
    items = _layout(lang)
    gaps = _solve_gaps(items)
    for item in items:
        if "gap" in item:
            _gap(doc, gaps[item["gap"]])
        else:
            _emit(doc, item)


def build(lang: str, out_docx: Path) -> None:
    doc = Document()
    md2gost._configure_styles(doc)
    md2gost._configure_page(doc)
    populate(doc, lang)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    print("[docx]", out_docx.name)


def latest_manuscript_date() -> str:
    """Date stamp of the newest rendered manuscript pair in defense/docs/.

    The title page carries no dated input — the stamp is a label only — but a
    pinned label is worse than no label: it is how the other four builders sent
    current content out under a June name, and it is why this one produced no
    August file at all while everything beside it was rebuilt.
    """
    docs = ROOT / "defense/docs"
    dates = sorted(
        m.group(1)
        for p in docs.glob("DISSERTATION_EN_GOST_*.docx")
        if (m := re.search(r"_(\d{4}-\d{2}-\d{2})\.docx$", p.name))
        and (docs / p.name.replace("_EN_", "_KZ_")).is_file()
    )
    if not dates:
        raise SystemExit(f"no DISSERTATION_{{EN,KZ}}_GOST_*.docx pair in {docs}")
    return dates[-1]


def main() -> None:
    ap = argparse.ArgumentParser(description="Build GOST TITLE PAGE (EN+KZ)")
    ap.add_argument("--date", default=None, help="output date stamp (default: newest manuscript)")
    ap.add_argument("--no-pdf", action="store_true")
    args = ap.parse_args()
    if args.date is None:
        args.date = latest_manuscript_date()
        print(f"[src ] newest manuscript: {args.date}")

    out_dir = ROOT / "defense/docs/front_matter"
    built = []
    for lang in ("en", "kz"):
        out = out_dir / f"TITLE_PAGE_{lang.upper()}_GOST_{args.date}.docx"
        build(lang, out)
        built.append(out)

    if not args.no_pdf:
        from docx2pdf import convert
        for out_docx in built:
            pdf = out_docx.with_suffix(".pdf")
            convert(str(out_docx), str(pdf))
            print("[pdf ]", pdf.name)


if __name__ == "__main__":
    main()
