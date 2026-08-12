"""Convert a council Markdown document into a GOST-formatted .docx (and optionally .pdf).

GOST parameters (per council/en/02-formatting/gost-formatting.md):
    A4, single line spacing, Times New Roman 14 pt,
    margins: left 30 mm, right 10 mm, top 20 mm, bottom 20 mm,
    page numbers centered at the bottom (not printed on the first page).

Markdown supported: # / ## / ### / #### headings, **bold**, *italic*, `code`,
numbered lists (1.), bullet lists (- / *), --- rule, blank-line paragraphs,
pipe tables, fenced code, LaTeX math, [FIG-x.x] markers, and ```mermaid
diagrams (rendered to images — see the Mermaid section below).

Usage:
    python md2gost.py INPUT.md [-o OUTPUT.docx] [--pdf]
"""

from __future__ import annotations

import argparse
import hashlib
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

FONT_NAME = "Times New Roman"
FONT_SIZE = 14  # pt
FIRST_LINE_INDENT_CM = 1.25

_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*].*?\*|`.+?`|\$[^$\n]+\$)")

# --- Version-marker scrubbing -------------------------------------------------
# Council deliverables are rendered OUTSIDE thesis/ (into defense/docs/). Per the
# project versioning policy, version markers — including the "V5" proper noun for
# the preprocessing pipeline — must never leak outside thesis/. The source .md in
# thesis/output/ is allowed to keep them; this converter strips them on the way
# out so the .docx/.pdf never carry a version. See PROJECT_MEMORY/strip-version-markers.md.

# Each pattern eats one adjacent space (leading where present) so removal leaves
# no double space and no space before punctuation — without touching whitespace
# elsewhere on the line (e.g. signature underscores spaced for layout).
# Parenthetical version tag, e.g. " (V5)", "(v5.1)", "(version 5.0)", "(нұсқа 5)".
_VER_PAREN = re.compile(
    r"[ \t]*\((?:[Vv][345](?:\.\d+)*|(?:version|версия|версии|нұсқа)[ \t]*[345](?:\.\d+)*)\)",
    re.IGNORECASE,
)
# Bare token, e.g. " V5", " v5.2", " V4.1", " V3" (leading space consumed).
_VER_TOKEN = re.compile(r"[ \t]*\b[Vv][345](?:\.\d+)*\b")
# Word form, e.g. " version 5.0", " версия 5", " нұсқа 5.1" (leading space consumed).
_VER_WORD = re.compile(
    r"[ \t]*\b(?:version|версия|версии|нұсқа)[ \t]*[345](?:\.\d+)*\b", re.IGNORECASE
)


def strip_version_markers(text: str) -> str:
    """Remove version markers (V3/V4/V5, decimals, word forms) from `text`.

    Council deliverables render outside thesis/, where no version marker —
    including the "V5" pipeline proper noun — may appear. Each pattern consumes
    the space preceding the marker so the surrounding text stays clean without
    rewriting unrelated whitespace.

    Args:
        text: Source Markdown that may legitimately contain version markers
            (it lives under thesis/).

    Returns:
        The text with version markers removed.
    """
    text = _VER_PAREN.sub("", text)
    text = _VER_TOKEN.sub("", text)
    text = _VER_WORD.sub("", text)
    return text


# --- Process-metadata scrubbing -----------------------------------------------
# Working source .md under thesis/ carries assembly/provenance banners as leading
# blockquotes ("> **Intermediate EN assembly — 2026-06-17.** …", "> **STAGE-G
# (final pass) — …**"), and experiment prose drafted from results/ can carry run
# dates ("прогон 2026-08-02"), artifact paths, and log references. None of that is
# dissertation content: it is internal process history, and it must not appear in
# the council deliverables rendered into defense/docs/. The source keeps it; this
# converter drops it on the way out — same contract as strip_version_markers.

# A blockquote line that is an assembly/provenance/status banner rather than a
# quotation. Matched on the banner vocabulary, so genuine block quotations in the
# text (which do not open with these markers) are left alone.
_PROC_BANNER = re.compile(
    r"^[ \t]*>.*\b(?:assembly|assembled|STAGE-[A-Z]|provenance|провенанс|"
    r"intermediate|промежуточн\w*|прогон\w*|NOT the final|черновик|"
    r"working draft|draft header|обновлено под)\b",
    re.IGNORECASE,
)
# Inline run-date reference, e.g. "(прогон 2026-08-02)", "прогон 02.08.2026",
# "run of 2026-08-02". Consumes the preceding space and an enclosing paren pair.
_RUN_DATE = re.compile(
    r"[ \t]*\(?(?:прогон\w*|run(?:\s+of)?)[ \t]*(?:от[ \t]*)?"
    r"\d{2,4}[-.]\d{2}[-.]\d{2,4}\)?",
    re.IGNORECASE,
)
# Bare artifact/log path reference, e.g. "experiments/outputs/exp1/summary.json",
# "VALUES.md", "predictions.npz", "*.log". Image extensions are deliberately
# excluded: figure markers legitimately point into experiments/outputs/ and are
# resolved by _FIG below, so stripping those paths would break figure rendering.
# The path body is built from dot/slash-separated segments so it can never end on
# a separator — otherwise a trailing sentence period would be swallowed with it.
_ARTIFACT_REF = re.compile(
    r"[ \t]*\(?(?:`?(?:experiments/)?outputs?(?:[/.][\w*-]+)+"
    r"(?<!\.png)(?<!\.jpg)(?<!\.jpeg)(?<!\.svg)(?<!\.gif)(?<!\.pdf)`?"
    r"|`?VALUES\.md`?|`?[\w-]+\.(?:log|npz|ckpt|pt)`?)\)?",
    re.IGNORECASE,
)
# A line carrying a figure marker is left untouched by artifact scrubbing.
_FIG_LINE = re.compile(r"\[FIG-[\w.]+:", re.IGNORECASE)


def strip_process_metadata(text: str) -> str:
    """Remove internal process history from text bound for defense/docs/.

    Drops assembly/provenance banner blockquotes wholesale, then removes inline
    run-date and raw-artifact references. Dissertation deliverables state results,
    not the run history that produced them; that history stays in results/ and
    PROJECT_MEMORY/ where it is needed for traceability.

    Args:
        text: Source Markdown that may legitimately carry process metadata
            (it lives under thesis/).

    Returns:
        The text with banner lines dropped and inline run/artifact references
        removed.
    """
    out = []
    for ln in text.splitlines():
        if _PROC_BANNER.match(ln):
            continue
        ln = _RUN_DATE.sub("", ln)
        if not _FIG_LINE.search(ln):  # figure markers keep their outputs/ paths
            ln = _ARTIFACT_REF.sub("", ln)
        out.append(ln)
    return "\n".join(out)


def _set_cell_font(run, *, bold=False, italic=False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    # Ensure the font also applies to complex-script / Cyrillic ranges.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)
    run.bold = bold
    run.italic = italic


# --- LaTeX math rendering -----------------------------------------------------
# The dissertation source carries math as LaTeX, inline (`$…$`) and display
# (`$$…$$`, optionally with `\tag{N}`). Word here has no equation engine, so this
# converter renders math as clean Unicode text with real super/subscript runs:
# `\beta\,A/L` → "βA/L", `$T/80$` → "T/80", `\frac{T}{80}` → "T/80",
# `p_t` → p with a subscript t, `A^k` → A with a superscript k. This removes the
# stray `$`, backslash commands and braces that otherwise leak into the .docx/.pdf.

_TEX_CMD = re.compile(r"\\([a-zA-Z]+)")

# Control symbols after a backslash (non-letter): spacing, escapes, line break.
_TEX_CTRL = {
    ",": "", "!": "", ";": " ", ":": " ", " ": " ", "\\": " ",
    "_": "_", "{": "{", "}": "}", "%": "%", "#": "#", "&": "&", "$": "$",
}

# Symbol commands → Unicode (Greek lower/upper, operators, relations, brackets).
_TEX_SYM = {
    # Greek (lowercase)
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ", "epsilon": "ε",
    "varepsilon": "ε", "zeta": "ζ", "eta": "η", "theta": "θ", "vartheta": "ϑ",
    "iota": "ι", "kappa": "κ", "lambda": "λ", "mu": "μ", "nu": "ν", "xi": "ξ",
    "pi": "π", "rho": "ρ", "sigma": "σ", "tau": "τ", "upsilon": "υ", "phi": "φ",
    "varphi": "φ", "chi": "χ", "psi": "ψ", "omega": "ω",
    # Greek (uppercase)
    "Gamma": "Γ", "Delta": "Δ", "Theta": "Θ", "Lambda": "Λ", "Xi": "Ξ",
    "Pi": "Π", "Sigma": "Σ", "Phi": "Φ", "Psi": "Ψ", "Omega": "Ω",
    # Operators / relations
    "cdot": "·", "times": "×", "div": "÷", "pm": "±", "mp": "∓", "ast": "∗",
    "ge": "≥", "geq": "≥", "le": "≤", "leq": "≤", "ne": "≠", "neq": "≠",
    "approx": "≈", "equiv": "≡", "sim": "∼", "propto": "∝",
    "cap": "∩", "cup": "∪", "subset": "⊂", "subseteq": "⊆", "supset": "⊃",
    "in": "∈", "notin": "∉", "forall": "∀", "exists": "∃",
    "sum": "Σ", "prod": "∏", "int": "∫", "partial": "∂", "nabla": "∇",
    "infty": "∞", "to": "→", "rightarrow": "→", "leftarrow": "←",
    "Rightarrow": "⇒", "leftrightarrow": "↔", "top": "⊤", "perp": "⊥",
    "angle": "∠", "lceil": "⌈", "rceil": "⌉", "lfloor": "⌊", "rfloor": "⌋",
    "cdots": "⋯", "ldots": "…", "dots": "…", "quad": "  ", "qquad": "    ",
}

# Operator/function names rendered upright (as the word itself).
_TEX_FUNC = {
    "min", "max", "log", "ln", "exp", "sin", "cos", "tan", "det", "lim",
    "deg", "gcd", "arg", "dim", "ker", "sup", "inf",
}

# One-argument commands whose content is rendered upright (drop the wrapper).
_TEX_WRAP = {
    "text", "mathrm", "mathbf", "mathit", "mathsf", "mathcal", "mathbb",
    "operatorname", "texttt", "boldsymbol", "mathtt", "textbf", "textit",
}

# Commands that render to nothing (delimiter sizing, style directives): the
# surrounding bracket/character they qualify is kept verbatim.
_TEX_DROP = {
    "left", "right", "big", "Big", "bigg", "Bigg", "bigl", "bigr",
    "Bigl", "Bigr", "biggl", "biggr", "Biggl", "Biggr", "displaystyle",
    "textstyle", "scriptstyle", "limits", "nolimits",
}


def _read_braces(s: str, i: int) -> tuple[str, int]:
    """Given s[i] == '{', return (inner_text, index_after_matching_'}')."""
    depth, j, n = 0, i, len(s)
    while j < n:
        if s[j] == "{":
            depth += 1
        elif s[j] == "}":
            depth -= 1
            if depth == 0:
                return s[i + 1:j], j + 1
        j += 1
    return s[i + 1:], n  # unbalanced: take the rest


def _read_script_arg(s: str, i: int) -> tuple[str, int]:
    """Read the argument of a `_`/`^` at index i: `{group}`, `\\command`, or one char."""
    n = len(s)
    if i >= n:
        return "", i
    if s[i] == "{":
        return _read_braces(s, i)
    if s[i] == "\\":
        m = _TEX_CMD.match(s, i)
        if m:
            return m.group(0), m.end()
        return s[i:i + 2], min(i + 2, n)
    return s[i], i + 1


def _tex_runs(s: str) -> list[tuple[str, str]]:
    """Parse a LaTeX math fragment into (text, script) runs.

    `script` is "" (baseline), "sub", or "sup". Commands, Greek letters and
    operators are mapped to Unicode; `\\frac{a}{b}` becomes "a/b"; sub/superscripts
    become tagged runs so the caller can apply real Word run formatting.
    """
    runs: list[tuple[str, str]] = []
    buf: list[str] = []
    n = len(s)

    def flush() -> None:
        if buf:
            runs.append(("".join(buf), ""))
            buf.clear()

    i = 0
    while i < n:
        c = s[i]
        if c == "\\":
            if i + 1 < n and not s[i + 1].isalpha():
                buf.append(_TEX_CTRL.get(s[i + 1], s[i + 1]))
                i += 2
                continue
            m = _TEX_CMD.match(s, i)
            name = m.group(1)
            i = m.end()
            if name in _TEX_DROP:
                pass  # sizing/style directive: render nothing, keep what follows
            elif name in _TEX_WRAP:
                inner, i = (_read_braces(s, i) if i < n and s[i] == "{" else ("", i))
                flush()
                runs.extend(_tex_runs(inner))
            elif name in ("frac", "dfrac", "tfrac"):
                a, b = "", ""
                if i < n and s[i] == "{":
                    a, i = _read_braces(s, i)
                while i < n and s[i] == " ":
                    i += 1
                if i < n and s[i] == "{":
                    b, i = _read_braces(s, i)
                flush()
                runs.extend(_frac_runs(a, b))
            elif name in _TEX_SYM:
                buf.append(_TEX_SYM[name])
            elif name in _TEX_FUNC:
                flush()
                runs.append((name, ""))
            else:
                buf.append(name)  # unknown: best-effort, drop the backslash
        elif c == "{":
            inner, i = _read_braces(s, i)
            flush()
            runs.extend(_tex_runs(inner))
        elif c == "}":
            i += 1  # stray closing brace
        elif c in "_^":
            script = "sub" if c == "_" else "sup"
            i += 1
            while i < n and s[i] == " ":
                i += 1
            arg, i = _read_script_arg(s, i)
            flush()
            runs.append((_tex_plain(arg), script))
        else:
            buf.append(c)
            i += 1
    flush()
    return runs


def _tex_plain(s: str) -> str:
    """Flatten a LaTeX fragment to plain text (used for script arguments/tags)."""
    return "".join(t for t, _ in _tex_runs(s))


def _frac_runs(a: str, b: str) -> list[tuple[str, str]]:
    """Render `\\frac{a}{b}` as a/b runs, parenthesising compound numerator/denominator."""
    return _maybe_paren(a) + [("/", "")] + _maybe_paren(b)


def _maybe_paren(group: str) -> list[tuple[str, str]]:
    runs = _tex_runs(group)
    plain = "".join(t for t, _ in runs).strip()
    if len(plain) > 1 and re.search(r"[+\-−·×/ ]", plain):
        return [("(", "")] + runs + [(")", "")]
    return runs


def _add_math_runs(paragraph, latex: str, *, bold=False, italic=False) -> None:
    """Render a LaTeX math fragment into `paragraph` as formatted runs."""
    for text, script in _tex_runs(latex):
        if not text:
            continue
        r = paragraph.add_run(text)
        _set_cell_font(r, bold=bold, italic=italic)
        if script == "sub":
            r.font.subscript = True
        elif script == "sup":
            r.font.superscript = True


def _split_tag(latex: str) -> tuple[str, str | None]:
    """Split a display equation's body from its `\\tag{…}` number, if any."""
    m = re.search(r"\\tag\{([^}]*)\}", latex)
    if not m:
        return latex, None
    return (latex[:m.start()] + latex[m.end():]).strip(), m.group(1)


def _add_equation(doc: Document, latex: str) -> None:
    """Render a display equation centred, with any `\\tag{}` number flush right (GOST)."""
    inner, tag = _split_tag(latex)
    usable = 170.0  # A4 text width: 210 − 30 (left) − 10 (right) mm
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(usable / 2), WD_TAB_ALIGNMENT.CENTER)
    if tag:
        p.paragraph_format.tab_stops.add_tab_stop(Mm(usable), WD_TAB_ALIGNMENT.RIGHT)
    _set_cell_font(p.add_run("\t"))
    _add_math_runs(p, inner.strip())
    if tag:
        _set_cell_font(p.add_run("\t(" + _tex_plain(tag).strip() + ")"))


def _add_runs(paragraph, text: str, *, bold=False, italic=False) -> None:
    """Add inline-formatted runs (**bold**, *italic*, `code`, `$math$`) to a paragraph."""
    for token in _INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            _set_cell_font(paragraph.add_run(token[2:-2]), bold=True, italic=italic)
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            _set_cell_font(r, bold=bold, italic=italic)
            r.font.name = "Consolas"
            r._element.get_or_add_rPr().find(qn("w:rFonts")).set(qn("w:ascii"), "Consolas")
            r._element.get_or_add_rPr().find(qn("w:rFonts")).set(qn("w:hAnsi"), "Consolas")
        elif token.startswith("$") and token.endswith("$") and len(token) >= 2:
            _add_math_runs(paragraph, token[1:-1], bold=bold, italic=italic)
        elif token.startswith("*") and token.endswith("*"):
            _set_cell_font(paragraph.add_run(token[1:-1]), bold=bold, italic=True)
        else:
            _set_cell_font(paragraph.add_run(token), bold=bold, italic=italic)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    pf = normal.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.different_first_page_header_footer = True  # no number on page 1


def _add_page_numbers(doc: Document) -> None:
    """Centered PAGE field in the footer; first-page footer left blank."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    _set_cell_font(run)
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)
    # Leave the first-page footer empty.
    section.first_page_footer.is_linked_to_previous = False


def _add_hrule(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_runs(p, text, bold=True)
        for r in p.runs:
            r.font.size = Pt(16)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_runs(p, text, bold=True, italic=(level >= 4))
    return p


def _body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(FIRST_LINE_INDENT_CM * 10)
    _add_runs(p, text)
    return p


def _list_item(doc: Document, marker: str, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Mm(12.5)
    p.paragraph_format.first_line_indent = Mm(-7.0)  # hanging
    _add_runs(p, f"{marker}\t{text}")
    return p


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Render a Markdown pipe-table as a bordered Word table (TNR, header bold).

    `rows` is the list of cell-text rows (the `|---|` separator already removed).
    Backward-compatible: only invoked when the source contains pipe tables, which
    council deliverables do not, so their rendering is unaffected.
    """
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ""  # clear default empty run
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Mm(0)
            p.paragraph_format.line_spacing = 1.0
            text = row[j] if j < len(row) else ""
            _add_runs(p, text, bold=(i == 0))
    # spacing paragraph after the table
    doc.add_paragraph()


def _add_code_block(doc: Document, code_lines: list[str]) -> None:
    """Render a fenced code block as left-aligned monospace lines (no indent)."""
    for ln in code_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Mm(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(FONT_SIZE - 2)
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), "Consolas")


_NUM = re.compile(r"^(\d+)\.\s+(.*)$")
_BUL = re.compile(r"^[-*]\s+(.*)$")
_HDR = re.compile(r"^(#{1,6})\s+(.*)$")
_TBL_ROW = re.compile(r"^\s*\|.*\|\s*$")
_TBL_SEP = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|?\s*$")
_FENCE = re.compile(r"^\s*```\s*([A-Za-z0-9_+-]*)")

# --- Mermaid diagram rendering ------------------------------------------------
# Appendix C supplies its four structural views (component, deployment, sequence,
# data) as Mermaid source and states that rendering to an image happens at
# conversion. A converter that treated the fence as an ordinary code block would
# deliver those views to the reader as monospace source, and the appendix would
# fail to discharge DIA-6.3. So a ```mermaid fence is rendered to a PNG and
# embedded, and a failure to render is loud rather than silent.
#
# Rendered PNGs are cached under defense/figures/mermaid/ keyed by a hash of the
# diagram source. The cache is a build input, not a scratch directory: the Kazakh
# Mermaid source is byte-identical to the English by design, so both editions hit
# the same entry, and a machine without Node can still build the document from the
# committed PNGs. Change the source and the key changes with it, so a stale image
# cannot survive an edit.

MERMAID_CACHE_DIR = "defense/figures/mermaid"

# Appendix C authors each diagram's caption as a bold line ABOVE its Mermaid
# fence. GOST places an illustration's caption under the illustration ("Слово
# «Рисунок» и его наименование помещают после пояснительных данных"), which is
# where every other caption in this document sits. The caption is therefore moved
# below its diagram at conversion time and set in the same centred form, so the
# appendix source stays readable and the output stays conventional.
_CAPTION_TOKEN = "\x01CAP\x01"
_DIAGRAM_CAPTION = re.compile(
    r"^\*\*(?:"
    r"(?P<en>(?:Diagram|Figure)\s+[\w.]+)"          # "Diagram C.1. Title."
    r"|(?P<kz>[\w.]+-(?:диаграмма|сурет))"          # "Б.1-диаграмма. Атауы."
    r")[.:]\s*(?P<title>.+?)\.?\*\*$"
)


def _relocate_diagram_captions(lines: list[str]) -> list[str]:
    """Move a bold diagram caption from above its Mermaid fence to below it."""
    out: list[str] = []
    i, n = 0, len(lines)
    while i < n:
        m = _DIAGRAM_CAPTION.match(lines[i].strip())
        if not m:
            out.append(lines[i])
            i += 1
            continue
        j = i + 1
        while j < n and not lines[j].strip():
            j += 1
        if j >= n or not lines[j].strip().startswith("```mermaid"):
            out.append(lines[i])          # not a diagram caption after all
            i += 1
            continue
        end = j + 1
        while end < n and not lines[end].strip().startswith("```"):
            end += 1
        if m.group("kz"):
            # "Б.1-диаграмма" -> "Диаграмма Б.1": the instruction's own example
            # puts the word first ("Рисунок 1 – Детали прибора"), and every other
            # caption this converter emits is label-first, so the two forms would
            # otherwise sit side by side in the same document.
            num, _, word = m.group("kz").rpartition("-")
            label = f"{word.capitalize()} {num}"
        else:
            label = m.group("en").strip()
        out.extend(lines[j:end + 1])
        out.append("")
        out.append(f"{_CAPTION_TOKEN}{label} – {m.group('title').strip()}")
        i = end + 1
    return out

# Chrome/Chromium locations tried when PUPPETEER_EXECUTABLE_PATH is unset. The
# candidate builds on several machines; reusing an installed browser avoids a
# per-machine Chromium download.
_CHROME_CANDIDATES = (
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    "/usr/bin/google-chrome",
    "/usr/bin/chromium",
    "/usr/bin/chromium-browser",
    "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
)

_mermaid_failures: list[str] = []


def _chrome_executable() -> str | None:
    """Path to an installed Chrome/Chromium, or None to let Puppeteer choose."""
    env = os.environ.get("PUPPETEER_EXECUTABLE_PATH")
    if env and Path(env).is_file():
        return env
    for cand in _CHROME_CANDIDATES:
        if Path(cand).is_file():
            return cand
    found = shutil.which("google-chrome") or shutil.which("chromium")
    return found


def _mmdc_command() -> list[str] | None:
    """Resolve the mermaid-cli invocation, or None if it cannot be found.

    Tried in order: the MMDC environment variable, a mermaid-cli installed into
    the repository's node_modules, one on PATH, then `npx` as a last resort
    (which needs network access on first use).
    """
    env = os.environ.get("MMDC")
    if env:
        return [env]
    here = Path(__file__).resolve().parent
    for base in (here, *here.parents):
        for name in ("mmdc.cmd", "mmdc"):
            cand = base / "node_modules" / ".bin" / name
            if cand.is_file():
                return [str(cand)]
    on_path = shutil.which("mmdc")
    if on_path:
        return [on_path]
    npx = shutil.which("npx")
    if npx:
        return [npx, "--yes", "@mermaid-js/mermaid-cli"]
    return None


def _render_mermaid(code: str, base_dir: Path) -> Path | None:
    """Render Mermaid `code` to a PNG and return its path, or None on failure.

    Uses the cached render when one exists, so the document builds without Node
    as long as the committed PNGs match the source.
    """
    digest = hashlib.sha256(code.strip().encode("utf-8")).hexdigest()[:12]
    cache_dir = base_dir / MERMAID_CACHE_DIR
    png = cache_dir / f"diagram_{digest}.png"
    if png.is_file():
        return png

    cmd = _mmdc_command()
    if cmd is None:
        _mermaid_failures.append(
            f"{digest}: mermaid-cli not found and no cached render at {png}"
        )
        return None

    cache_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "diagram.mmd"
        src.write_text(code.strip() + "\n", encoding="utf-8")
        argv = [*cmd, "-i", str(src), "-o", str(png), "-b", "white", "-s", "3"]
        chrome = _chrome_executable()
        if chrome:
            cfg = Path(tmp) / "puppeteer.json"
            cfg.write_text(
                '{"executablePath": %s, "args": ["--no-sandbox", '
                '"--disable-setuid-sandbox"]}' % _json_str(chrome),
                encoding="utf-8",
            )
            argv += ["-p", str(cfg)]
        try:
            proc = subprocess.run(argv, capture_output=True, text=True, timeout=180)
        except (OSError, subprocess.SubprocessError) as exc:
            _mermaid_failures.append(f"{digest}: {exc}")
            return None
    if proc.returncode != 0 or not png.is_file():
        detail = (proc.stderr or proc.stdout or "").strip().splitlines()
        _mermaid_failures.append(f"{digest}: {detail[-1] if detail else 'render failed'}")
        return None
    return png


def _json_str(s: str) -> str:
    """Minimal JSON string literal (avoids importing json for one value)."""
    return '"' + s.replace("\\", "\\\\").replace('"', '\\"') + '"'


def mermaid_failures() -> list[str]:
    """Diagrams that failed to render during the conversions run so far."""
    return list(_mermaid_failures)


def _add_mermaid(doc: Document, code_lines: list[str], base_dir: Path) -> None:
    """Embed a Mermaid diagram as a centred image; fall back to source on failure.

    The fallback is deliberately visible — the caller reports the failure and
    exits non-zero — because a diagram silently shipped as source is the defect
    this function exists to prevent.
    """
    img = _render_mermaid("\n".join(code_lines), base_dir)
    if img is None:
        _add_code_block(doc, code_lines)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(str(img), width=Mm(_fit_width_mm(img)))

# --- Figure placeholders ------------------------------------------------------
# Drafts carry assets as inline text markers `[FIG-3.1: caption — path/img.png]`.
# This converter resolves each to an embedded image with a GOST caption below it
# ("Figure N – Title" / KZ "Сурет N – Атауы"), replacing the inline marker with a
# cross-reference. `pre`("…in ") / `post`(" …суретінде") are absorbed so the
# reference reads naturally in either language.
#
# Four marker prefixes occur, and they are NOT interchangeable:
#   FIG / FIGURE — a figure, numbered within its chapter or appendix (E.1…E.54).
#   APP          — an appendix exhibit carrying only the appendix letter, so the
#                  sequence number is assigned here in order of appearance.
#   DIA          — a diagram. It takes its own caption series because DIA-6.1 and
#                  FIG-6.1 both exist: labelling both "Figure 6.1" would put two
#                  different images under one number.
#   TAB          — a table caption. Never an image: the table itself follows as
#                  Markdown, and GOST places its caption above it, which is where
#                  the marker already sits.
# Numbers may carry an appendix letter (E.1, D), so the pattern is not digits-only
# — matching digits alone left all 54 Appendix-E plates in the document as raw
# bracket text with their file paths showing.
_FIG = re.compile(
    r"(?P<pre>\b[Ii]n\s+)?"
    r"\[(?P<kind>FIG|FIGURE|APP|DIA|TAB)-(?P<num>[A-Za-z]?[0-9]*(?:\.[0-9]+)?):"
    r"\s*(?P<body>[^\]]*)\]"
    r"(?P<post>\s+сурет\w*)?"
)

# Caption word per marker kind and language.
_LABELS = {
    "en": {"FIG": "Figure", "FIGURE": "Figure", "APP": "Figure",
           "DIA": "Diagram", "TAB": "Table"},
    "kz": {"FIG": "Сурет", "FIGURE": "Сурет", "APP": "Сурет",
           "DIA": "Диаграмма", "TAB": "Кесте"},
}
_DASH = re.compile(r"\s+[—–-]\s+")  # caption — target separator (em/en/hyphen)
# A marker wrapped in backticks is still a marker, not a code span.
_FIG_TICKS = re.compile(r"`(\[(?:FIG|FIGURE|APP|DIA|TAB)-[^\]]*\])`")


def _png_size(p: Path):
    """Return (w, h) in pixels for a PNG, else None."""
    try:
        with open(p, "rb") as f:
            head = f.read(26)
        if head[:8] == b"\x89PNG\r\n\x1a\n" and head[12:16] == b"IHDR":
            return int.from_bytes(head[16:20], "big"), int.from_bytes(head[20:24], "big")
    except OSError:
        pass
    return None


def _fit_width_mm(p: Path, maxw: float = 165.0, maxh: float = 215.0) -> float:
    """Width (mm) that fits the image inside the text box, preserving aspect."""
    sz = _png_size(p)
    if not sz:
        return maxw * 0.9
    w, h = sz
    return min(maxw, maxh * w / h)


def _parse_fig_body(body: str, base: Path):
    """Split '[caption — target]' into (caption, resolved_image_path_or_None)."""
    segs = _DASH.split(body)
    if len(segs) >= 2:
        caption = " – ".join(s.strip() for s in segs[:-1]).strip()
        target = segs[-1].strip()
    else:
        caption, target = body.strip(), ""
    target = target.strip(" `").replace("/…/", "/").replace("…/", "").replace("/…", "")
    img = None
    if re.search(r"\.(png|jpe?g)$", target, re.I):
        cand = (base / target).resolve()
        if cand.is_file():
            img = cand
    return caption, img


# Images are embedded at the resolution the page can actually print. The 54
# Appendix-E plates are 2954 px wide and sit 165 mm wide — 455 dpi, of which a
# printer uses none above ~300 — and embedding them at native size put 73 MB of
# plates into an 86 MB document. Downscaled copies are cached beside the output.
PRINT_DPI = 300
PRINT_CACHE_DIR = "defense/docs/.print_cache"


def _print_ready(img: Path, width_mm: float, base_dir: Path) -> Path:
    """A copy of `img` downscaled to PRINT_DPI at `width_mm`, or `img` itself.

    Returns the original when it is already at or below the print resolution, or
    when Pillow is unavailable — an oversized image is a size problem, never a
    correctness one, so it is never worth failing the build over.
    """
    target = int(width_mm / 25.4 * PRINT_DPI)
    sz = _png_size(img)
    if not sz or sz[0] <= target * 1.1:
        return img
    try:
        from PIL import Image
    except ImportError:
        return img
    st = img.stat()
    key = hashlib.sha256(
        f"{img}|{st.st_mtime_ns}|{st.st_size}|{target}".encode("utf-8")
    ).hexdigest()[:12]
    cache = base_dir / PRINT_CACHE_DIR
    hit = next(cache.glob(f"{img.stem}_{key}.*"), None) if cache.is_dir() else None
    if hit:
        return hit
    cache.mkdir(parents=True, exist_ok=True)
    # Encode both ways and let the measured sizes choose, rather than guessing
    # from the content. PNG is kept unless JPEG is at least twice as small: on
    # the photographic plates JPEG wins by ~7x and is worth taking, while on
    # line art it saves ~50 KB and is not worth the ringing around type.
    try:
        with Image.open(img) as im:
            im = im.convert("RGB")
            small = im.resize((target, round(sz[1] * target / sz[0])), Image.LANCZOS)
            png, jpg = cache / f"{img.stem}_{key}.png", cache / f"{img.stem}_{key}.jpg"
            small.save(png, "PNG", optimize=True)
            small.save(jpg, "JPEG", quality=90, optimize=True)
    except OSError:
        return img
    keep, drop = (jpg, png) if jpg.stat().st_size * 2 <= png.stat().st_size else (png, jpg)
    drop.unlink(missing_ok=True)
    return keep


def _insert_figure(doc: Document, label: str, num: str, caption: str, img: Path | None,
                   *, note_missing: bool = True, base_dir: Path | None = None):
    if img is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Mm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.keep_with_next = True
        w = _fit_width_mm(img)
        src = _print_ready(img, w, base_dir) if base_dir is not None else img
        p.add_run().add_picture(str(src), width=Mm(w))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Mm(0)
    c.paragraph_format.space_after = Pt(6)
    text = f"{label} {num} – {caption}" if num else f"{label} – {caption}"
    if img is None and note_missing:
        text += " [ресурс дайындалуда]" if label == "Сурет" else " [asset to be created]"
    _add_runs(c, text)


def _parse_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def render_into(
    doc,
    text: str,
    *,
    lang: str = "en",
    base_dir: Path | None = None,
) -> None:
    """Render Markdown `text` into an existing (already-configured) document.

    Contains the full Markdown parsing loop but performs no page setup, footer,
    or save — so callers can compose several Markdown bodies into one document
    (e.g. the front-matter + manuscript bundle). `convert()` wraps this for the
    single-file case. Version-marker scrubbing is the caller's responsibility
    here (convert() still does it).
    """
    labels = _LABELS["kz" if lang == "kz" else "en"]
    if base_dir is None:
        base_dir = Path(".")
    lines = _relocate_diagram_captions(text.splitlines())

    figs: dict[tuple, dict] = {}   # (kind, num, target) -> registration
    seq: dict[tuple, int] = {}     # (kind, num) -> next sequence for letter-only ids

    def _key(m: re.Match) -> tuple:
        return (m.group("kind").upper(), m.group("num"), m.group("body").strip())

    def _register(m: re.Match) -> dict:
        """Resolve a marker once, assigning its caption number and image."""
        k = _key(m)
        if k not in figs:
            kind, num = k[0], k[1]
            caption, img = _parse_fig_body(m.group("body"), base_dir)
            if "." not in num:
                # An appendix-letter id (APP-D) carries no sequence of its own;
                # number the exhibits in order of appearance: D.1, D.2, …
                seq[(kind, num)] = seq.get((kind, num), 0) + 1
                num = f"{num}.{seq[(kind, num)]}" if num else str(seq[(kind, num)])
            figs[k] = {"label": labels.get(k[0], labels["FIG"]), "num": num,
                       "caption": caption, "img": img, "placed": False}
        return figs[k]

    def _fig_inline(m: re.Match) -> str:
        """Register the marker and return the text that replaces it in the prose."""
        f = _register(m)
        if f["img"] is None:
            # A marker with no image is a cross-reference, not a figure: the
            # sentence around it already names the target ("…given in Appendix C"),
            # so the bracket is dropped rather than printed at the reader.
            return ""
        if lang == "kz":
            post = (m.group("post") or "").strip()
            return f"{f['num']}-{post}" if post else f"({f['num']}-сурет)"
        pre = m.group("pre") or ""
        return f"{pre}{f['label']} {f['num']}" if pre else f"({f['label']} {f['num']})"

    def _emit(f: dict, *, note_missing: bool = True) -> None:
        if f["placed"]:
            return
        _insert_figure(doc, f["label"], f["num"], f["caption"], f["img"],
                       note_missing=note_missing, base_dir=base_dir)
        f["placed"] = True

    def resolve_markers(raw: str) -> tuple[str, list[dict], bool]:
        """Substitute asset markers in `raw`.

        Returns the prose with markers replaced by cross-references, the assets
        to emit after it, and whether the text was nothing but markers (a
        standalone caption line, which becomes a figure block on its own).
        """
        raw = _FIG_TICKS.sub(r"\1", raw)  # a marker wrapped in `…` is not code
        found = list(_FIG.finditer(raw))
        if not found:
            return raw, [], False
        if not _FIG.sub("", raw).strip():
            return "", [_register(m) for m in found], True
        regs = [_register(m) for m in found]
        cleaned = re.sub(r"[ \t]{2,}", " ", _FIG.sub(_fig_inline, raw)).strip()
        cleaned = re.sub(r"\s+([,.;:])", r"\1", cleaned)
        return cleaned, regs, False

    def emit_after(regs: list[dict]) -> None:
        for f in regs:
            if f["img"] is not None:
                _emit(f)

    buf: list[str] = []

    def flush_paragraph() -> None:
        if not buf:
            return
        raw = " ".join(buf).strip()
        buf.clear()
        cleaned, regs, standalone = resolve_markers(raw)
        if standalone:
            # Nothing but markers — the Appendix-E plate list, a table caption.
            # A table caption stands alone above its Markdown table, so its
            # absent image is the normal case rather than a missing asset.
            for f in regs:
                _emit(f, note_missing=f["label"] != labels["TAB"])
            return
        _body(doc, cleaned)
        emit_after(regs)

    tbl_buf: list[list[str]] = []

    def flush_table() -> None:
        if tbl_buf:
            _add_table(doc, tbl_buf)
            tbl_buf.clear()

    in_code = False
    code_lang = ""
    code_buf: list[str] = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        # fenced block: collect raw lines verbatim until the closing fence. A
        # ```mermaid fence is a diagram, rendered to an image rather than set as
        # source (Appendix C states rendering happens at conversion time).
        m_fence = _FENCE.match(stripped)
        if m_fence:
            if in_code:
                if code_lang.lower() == "mermaid":
                    _add_mermaid(doc, code_buf, base_dir)
                else:
                    _add_code_block(doc, code_buf)
                code_buf.clear()
                code_lang = ""
                in_code = False
            else:
                flush_paragraph()
                flush_table()
                code_lang = m_fence.group(1) or ""
                in_code = True
            continue
        if in_code:
            code_buf.append(raw.rstrip("\n"))
            continue

        # pipe table: accumulate consecutive table rows, skip the |---| separator
        if _TBL_ROW.match(line):
            flush_paragraph()
            if _TBL_SEP.match(line):
                continue
            tbl_buf.append(_parse_table_row(line))
            continue
        else:
            flush_table()

        if stripped.startswith(_CAPTION_TOKEN):
            flush_paragraph()
            c = doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.first_line_indent = Mm(0)
            c.paragraph_format.space_after = Pt(6)
            _add_runs(c, stripped[len(_CAPTION_TOKEN):])
            continue

        if not stripped:
            flush_paragraph()
            continue
        if stripped == "---" or set(stripped) == {"-"} and len(stripped) >= 3:
            flush_paragraph()
            _add_hrule(doc)
            continue

        # display equation on its own line: $$ … $$  (optionally with \tag{N})
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) >= 4:
            flush_paragraph()
            flush_table()
            _add_equation(doc, stripped[2:-2])
            continue

        m = _HDR.match(stripped)
        if m:
            flush_paragraph()
            _heading(doc, m.group(2).strip(), len(m.group(1)))
            continue

        # List items carry asset markers too — the Appendix-E plate list and the
        # Appendix-D publication confirmations are bulleted. Resolving markers
        # only in paragraphs left all sixty of those in the document as raw
        # bracket text with their file paths showing.
        m = _NUM.match(stripped) or _BUL.match(stripped)
        if m:
            flush_paragraph()
            marker = f"{m.group(1)}." if m.re is _NUM else "•"
            item = (m.group(2) if m.re is _NUM else m.group(1)).strip()
            cleaned, regs, standalone = resolve_markers(item)
            if standalone:
                for f in regs:
                    _emit(f, note_missing=f["label"] != labels["TAB"])
            else:
                _list_item(doc, marker, cleaned)
                emit_after(regs)
            continue

        buf.append(stripped)

    if in_code and code_buf:
        if code_lang.lower() == "mermaid":
            _add_mermaid(doc, code_buf, base_dir)
        else:
            _add_code_block(doc, code_buf)
    flush_table()
    flush_paragraph()


def convert(
    md_path: Path,
    docx_path: Path,
    *,
    strip_versions: bool = True,
    strip_process: bool = True,
    lang: str | None = None,
    base_dir: Path | None = None,
) -> None:
    text = md_path.read_text(encoding="utf-8")
    if strip_versions:
        text = strip_version_markers(text)
    if strip_process:
        text = strip_process_metadata(text)
    if lang is None:
        lang = "kz" if "_KZ_" in md_path.name else "en"
    if base_dir is None:  # repo root: walk up until a dir containing defense/
        base_dir = md_path.resolve().parent
        while base_dir.parent != base_dir and not (base_dir / "defense").is_dir():
            base_dir = base_dir.parent
    doc = Document()
    _configure_styles(doc)
    _configure_page(doc)
    render_into(doc, text, lang=lang, base_dir=base_dir)
    _add_page_numbers(doc)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> None:
    ap = argparse.ArgumentParser(description="Markdown -> GOST .docx (+ optional .pdf)")
    ap.add_argument("input", type=Path, help="input .md file")
    ap.add_argument("-o", "--output", type=Path, help="output .docx (default: alongside input)")
    ap.add_argument("--pdf", action="store_true", help="also render a .pdf via MS Word")
    args = ap.parse_args()

    md_path: Path = args.input
    docx_path: Path = args.output or md_path.with_suffix(".docx")
    convert(md_path, docx_path)
    print(f"[docx] {docx_path}")

    if _mermaid_failures:
        # A diagram that reached the reader as source would be a defect in the
        # document, not a warning about the build — so say so and fail.
        print(
            f"[FAIL] {len(_mermaid_failures)} Mermaid diagram(s) shipped as source:",
            file=sys.stderr,
        )
        for f in _mermaid_failures:
            print(f"       {f}", file=sys.stderr)
        raise SystemExit(1)

    if args.pdf:
        from docx2pdf import convert as to_pdf

        pdf_path = docx_path.with_suffix(".pdf")
        to_pdf(str(docx_path), str(pdf_path))
        print(f"[pdf ] {pdf_path}")


if __name__ == "__main__":
    main()
